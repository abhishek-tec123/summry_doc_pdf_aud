# # # with audio file----------------------------------------------------------------------------------------------------------------



# from flask import Flask, request, jsonify
# import os
# import time
# import io
# import logging
# import requests
# import PyPDF2
# import unicodedata
# import google.generativeai as genai
# from dotenv import load_dotenv
# from google.api_core.exceptions import ResourceExhausted
# from docx import Document

# # Configure logging
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# # Initialize Flask app
# app = Flask(__name__)

# # Load environment variables
# load_dotenv()
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     logging.error("API key not found. Please set the GOOGLE_API_KEY environment variable.")
#     exit(1)
# genai.configure(api_key=api_key)

# def extract_text_from_pdf(pdf_stream: io.BytesIO) -> str:
#     """Extract text from a PDF stream."""
#     logging.info("Extracting text from the PDF.")
#     pdf_reader = PyPDF2.PdfReader(pdf_stream)
#     text = ""
#     for page in pdf_reader.pages:
#         page_text = page.extract_text() or ""
#         text += page_text
#     return text

# def normalize_text(text: str) -> str:
#     """Normalize text to handle special characters."""
#     return unicodedata.normalize('NFKC', text)

# def chunk_text(text: str, chunk_size: int = 5000) -> list:
#     """Chunk text into smaller pieces."""
#     text = normalize_text(text)
#     chunks = []
#     start = 0
#     while start < len(text):
#         end = start + chunk_size
#         # Ensure we do not split surrogate pairs
#         if end < len(text):
#             while end < len(text) and text[end] in ('\uD800', '\uDBFF'):
#                 end += 1
#         chunks.append(text[start:end])
#         start = end
#     return chunks

# def summarize_chunk(llm, chunk: str, retries: int = 3, delay: int = 1) -> str:
#     """Summarize a text chunk using a language model."""
#     prompt = f"Please provide a well-defined summary of the following text:\n TEXT: {chunk}"
#     for attempt in range(retries):
#         try:
#             response = llm.generate_content(prompt)
#             # Log the response type and content
#             # logging.info(f"Response type: {type(response)}")
#             # logging.info(f"Response content: {response}")
#             # Assuming the response object has a 'text' attribute or method
#             summary_text = response.text  # Adjust this based on actual response object
#             return summary_text
#         except ResourceExhausted as e:
#             logging.warning(f"Resource exhausted error: {e}. Retrying in {delay} seconds...")
#             time.sleep(delay)
#             delay *= 2  # Exponential backoff
#         except Exception as e:
#             logging.error(f"Unexpected error: {e}. Retrying in {delay} seconds...")
#             time.sleep(delay)
#             delay *= 2  # Exponential backoff
#     raise Exception("Max retries exceeded")

# def extract_text_from_docx(docx_stream: io.BytesIO) -> str:
#     """Extract text from a DOCX stream."""
#     logging.info("Extracting text from the DOCX.")
#     doc = Document(docx_stream)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + '\n'
#     return text

# def extract_text_from_txt(txt_stream: io.BytesIO) -> str:
#     """Extract text from a TXT stream."""
#     logging.info("Extracting text from the TXT.")
#     return txt_stream.read().decode('utf-8')

# def handle_file(file_stream: io.BytesIO, file_type: str) -> str:
#     """Handle different types of files."""
#     if file_type == 'pdf':
#         return extract_text_from_pdf(file_stream)
#     elif file_type == 'docx':
#         return extract_text_from_docx(file_stream)
#     elif file_type == 'txt':
#         return extract_text_from_txt(file_stream)
#     else:
#         raise ValueError("Unsupported file format")

# def extract_transcript(audio_file_path):
#     """Extract transcript from an audio file using generative AI model."""
#     # Initialize the generative model
#     model = genai.GenerativeModel('gemini-1.5-flash')
    
#     # Upload the audio file
#     audio = genai.upload_file(path=audio_file_path)
    
#     # Generate content with a prompt to extract the transcript
#     response = model.generate_content(['extract well structured transcript from audio', audio])
    
#     # Return the transcript text
#     return response.text

# @app.route('/upload/pdf', methods=['POST'])
# def upload_file():
#     """Handle file upload from URL or local file."""
#     if 'file' in request.files:
#         # Handle file upload from local system
#         file = request.files['file']
#         if file.filename == '':
#             return jsonify({'error': 'No selected file'}), 400
#         file_type = file.filename.rsplit('.', 1)[-1].lower()
#         if file_type not in ['pdf', 'docx', 'txt']:
#             return jsonify({'error': 'Unsupported file format'}), 400
#         pdf_stream = io.BytesIO(file.read())
#     elif 'url' in request.json:
#         # Handle PDF from URL
#         url = request.json['url']
#         if not url:
#             return jsonify({'error': 'No URL provided'}), 400
#         try:
#             response = requests.get(url, stream=True)
#             response.raise_for_status()  # Ensure we notice bad responses
#             content_type = response.headers.get('Content-Type', '')
#             if 'application/pdf' in content_type:
#                 file_type = 'pdf'
#             elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
#                 file_type = 'docx'
#             elif 'text/plain' in content_type:
#                 file_type = 'txt'
#             else:
#                 return jsonify({'error': 'Unsupported file format from URL'}), 400
#             pdf_stream = io.BytesIO(response.content)
#         except requests.RequestException as e:
#             logging.error(f"Error fetching file from URL: {e}")
#             return jsonify({'error': 'Error fetching file from URL'}), 500
#     else:
#         return jsonify({'error': 'No file or URL provided'}), 400

#     try:
#         text = handle_file(pdf_stream, file_type)
#         chunks = chunk_text(text)
        
#         llm = genai.GenerativeModel("gemini-pro")
#         summaries = []

#         logging.info("Starting the summarization process.")
#         for i, chunk in enumerate(chunks):
#             try:
#                 summary = summarize_chunk(llm, chunk)
#                 if summary:  # Only append if summarization was successful
#                     summaries.append(summary)
#                     logging.info(f"Chunk {i+1}/{len(chunks)} summarized successfully.")
#                 else:
#                     logging.warning(f"Chunk {i+1} could not be summarized. Skipping.")
#                 time.sleep(1)  # Add delay between requests to avoid rate limit
#             except Exception as e:
#                 logging.error(f"Error processing chunk {i+1}: {e}")
#                 continue

#         # Ensure no None values are in summaries
#         final_summary = " ".join(filter(None, summaries))
#         return jsonify({'summary': final_summary})
#     except Exception as e:
#         logging.error(f"An error occurred: {e}")
#         return jsonify({'error': str(e)}), 500

# @app.route('/upload/audio', methods=['POST'])
# def upload_audio():
#     """Handle audio file upload and extract transcript."""
#     if 'file' in request.files:
#         # Handle file upload from local system
#         file = request.files['file']
#         if file.filename == '':
#             return jsonify({'error': 'No selected file'}), 400
#         file_type = file.filename.rsplit('.', 1)[-1].lower()
#         if file_type not in ['wav', 'mp3', 'flac']:
#             return jsonify({'error': 'Unsupported audio format'}), 400
#         audio_path = f"/tmp/{file.filename}"
#         file.save(audio_path)
#     elif 'url' in request.json:
#         # Handle audio from URL
#         url = request.json['url']
#         if not url:
#             return jsonify({'error': 'No URL provided'}), 400
#         try:
#             response = requests.get(url, stream=True)
#             response.raise_for_status()  # Ensure we notice bad responses
#             content_type = response.headers.get('Content-Type', '')
#             if 'audio' not in content_type:
#                 return jsonify({'error': 'Unsupported audio format from URL'}), 400
#             audio_path = f"/tmp/{url.split('/')[-1]}"
#             with open(audio_path, 'wb') as f:
#                 f.write(response.content)
#         except requests.RequestException as e:
#             logging.error(f"Error fetching audio from URL: {e}")
#             return jsonify({'error': 'Error fetching audio from URL'}), 500
#     else:
#         return jsonify({'error': 'No file or URL provided'}), 400

#     try:
#         transcript = extract_transcript(audio_path)
#         return jsonify({'transcript': transcript})
#     except Exception as e:
#         logging.error(f"An error occurred during transcript extraction: {e}")
#         return jsonify({'error': str(e)}), 500

# if __name__ == "__main__":
#     app.run(debug=True)


# with video÷\-----------------------------------------------------------------------------------------------------------------------------------


from flask import Flask, request, jsonify
import os
import time
import io
import logging
import requests
import PyPDF2
import unicodedata
import google.generativeai as genai
from dotenv import load_dotenv
from google.api_core.exceptions import ResourceExhausted
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Initialize Flask app
app = Flask(__name__)

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    logging.error("API key not found. Please set the GOOGLE_API_KEY environment variable.")
    exit(1)
genai.configure(api_key=api_key)

def extract_text_from_pdf(pdf_stream: io.BytesIO) -> str:
    """Extract text from a PDF stream."""
    logging.info("Extracting text from the PDF.")
    pdf_reader = PyPDF2.PdfReader(pdf_stream)
    text = ""
    for page in pdf_reader.pages:
        page_text = page.extract_text() or ""
        text += page_text
    return text

def normalize_text(text: str) -> str:
    """Normalize text to handle special characters."""
    return unicodedata.normalize('NFKC', text)

def chunk_text(text: str, chunk_size: int = 5000) -> list:
    """Chunk text into smaller pieces."""
    text = normalize_text(text)
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        # Ensure we do not split surrogate pairs
        if end < len(text):
            while end < len(text) and text[end] in ('\uD800', '\uDBFF'):
                end += 1
        chunks.append(text[start:end])
        start = end
    return chunks

def summarize_chunk(llm, chunk: str, retries: int = 3, delay: int = 1) -> str:
    """Summarize a text chunk using a language model."""
    prompt = f"Please provide a well-defined summary of the following text:\n TEXT: {chunk}"
    for attempt in range(retries):
        try:
            response = llm.generate_content(prompt)
            summary_text = response.text  # Adjust this based on actual response object
            return summary_text
        except ResourceExhausted as e:
            logging.warning(f"Resource exhausted error: {e}. Retrying in {delay} seconds...")
            time.sleep(delay)
            delay *= 2  # Exponential backoff
        except Exception as e:
            logging.error(f"Unexpected error: {e}. Retrying in {delay} seconds...")
            time.sleep(delay)
            delay *= 2  # Exponential backoff
    raise Exception("Max retries exceeded")

def extract_text_from_docx(docx_stream: io.BytesIO) -> str:
    """Extract text from a DOCX stream."""
    logging.info("Extracting text from the DOCX.")
    doc = Document(docx_stream)
    text = ""
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

def extract_text_from_txt(txt_stream: io.BytesIO) -> str:
    """Extract text from a TXT stream."""
    logging.info("Extracting text from the TXT.")
    return txt_stream.read().decode('utf-8')

def handle_file(file_stream: io.BytesIO, file_type: str) -> str:
    """Handle different types of files."""
    if file_type == 'pdf':
        return extract_text_from_pdf(file_stream)
    elif file_type == 'docx':
        return extract_text_from_docx(file_stream)
    elif file_type == 'txt':
        return extract_text_from_txt(file_stream)
    else:
        raise ValueError("Unsupported file format")

def extract_transcript(audio_file_path):
    """Extract transcript from an audio file using generative AI model."""
    model = genai.GenerativeModel('gemini-1.5-flash')
    audio = genai.upload_file(path=audio_file_path)
    response = model.generate_content(['extract well structured transcript from audio', audio])
    return response.text

def get_transcript_from_video(video_path):
    """Extract transcript from a video file using generative AI model."""
    try:
        # Upload the video file
        video = genai.upload_file(path=video_path)
        
        # Check the status of the upload
        while video.state.name != 'ACTIVE':
            print('.', end="")
            time.sleep(10)
            video = genai.get_file(video.name)
            
            if video.state.name not in ['processing', 'ACTIVE']:
                raise Exception(f"Unexpected video state: {video.state.name}")

        print('\nVideo uploaded successfully and is in ACTIVE state')

        # Initialize the generative model
        model = genai.GenerativeModel('gemini-1.5-flash')

        # Generate the content using the model
        prompt = "write transcript of video"
        response = model.generate_content([prompt, video])
        
        return response.text

    except Exception as e:
        print(f"An error occurred: {e}")
        return None

@app.route('/upload/pdf', methods=['POST'])
def upload_file():
    """Handle file upload from URL or local file."""
    if 'file' in request.files:
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        file_type = file.filename.rsplit('.', 1)[-1].lower()
        if file_type not in ['pdf', 'docx', 'txt']:
            return jsonify({'error': 'Unsupported file format'}), 400
        pdf_stream = io.BytesIO(file.read())
    elif 'url' in request.json:
        url = request.json['url']
        if not url:
            return jsonify({'error': 'No URL provided'}), 400
        try:
            response = requests.get(url, stream=True)
            response.raise_for_status()
            content_type = response.headers.get('Content-Type', '')
            if 'application/pdf' in content_type:
                file_type = 'pdf'
            elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
                file_type = 'docx'
            elif 'text/plain' in content_type:
                file_type = 'txt'
            else:
                return jsonify({'error': 'Unsupported file format from URL'}), 400
            pdf_stream = io.BytesIO(response.content)
        except requests.RequestException as e:
            logging.error(f"Error fetching file from URL: {e}")
            return jsonify({'error': 'Error fetching file from URL'}), 500
    else:
        return jsonify({'error': 'No file or URL provided'}), 400

    try:
        text = handle_file(pdf_stream, file_type)
        chunks = chunk_text(text)
        
        llm = genai.GenerativeModel("gemini-pro")
        summaries = []

        logging.info("Starting the summarization process.")
        for i, chunk in enumerate(chunks):
            try:
                summary = summarize_chunk(llm, chunk)
                if summary:
                    summaries.append(summary)
                    logging.info(f"Chunk {i+1}/{len(chunks)} summarized successfully.")
                else:
                    logging.warning(f"Chunk {i+1} could not be summarized. Skipping.")
                time.sleep(1)
            except Exception as e:
                logging.error(f"Error processing chunk {i+1}: {e}")
                continue

        final_summary = " ".join(filter(None, summaries))
        return jsonify({'summary': final_summary})
    except Exception as e:
        logging.error(f"An error occurred: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/upload/audio', methods=['POST'])
def upload_audio():
    """Handle audio file upload and extract transcript."""
    if 'file' in request.files:
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        file_type = file.filename.rsplit('.', 1)[-1].lower()
        if file_type not in ['wav', 'mp3', 'flac']:
            return jsonify({'error': 'Unsupported audio format'}), 400
        audio_path = f"/tmp/{file.filename}"
        file.save(audio_path)
    elif 'url' in request.json:
        url = request.json['url']
        if not url:
            return jsonify({'error': 'No URL provided'}), 400
        try:
            response = requests.get(url, stream=True)
            response.raise_for_status()
            content_type = response.headers.get('Content-Type', '')
            if 'audio' not in content_type:
                return jsonify({'error': 'Unsupported audio format from URL'}), 400
            audio_path = f"/tmp/{url.split('/')[-1]}"
            with open(audio_path, 'wb') as f:
                f.write(response.content)
        except requests.RequestException as e:
            logging.error(f"Error fetching audio from URL: {e}")
            return jsonify({'error': 'Error fetching audio from URL'}), 500
    else:
        return jsonify({'error': 'No file or URL provided'}), 400

    try:
        transcript = extract_transcript(audio_path)
        return jsonify({'transcript': transcript})
    except Exception as e:
        logging.error(f"An error occurred during transcript extraction: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/upload/video', methods=['POST'])
def upload_video():
    """Handle video file upload and extract transcript."""
    if 'file' in request.files:
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        file_type = file.filename.rsplit('.', 1)[-1].lower()
        if file_type not in ['mp4', 'avi', 'mov']:
            return jsonify({'error': 'Unsupported video format'}), 400
        video_path = f"/tmp/{file.filename}"
        file.save(video_path)
    elif 'url' in request.json:
        url = request.json['url']
        if not url:
            return jsonify({'error': 'No URL provided'}), 400
        try:
            response = requests.get(url, stream=True)
            response.raise_for_status()
            content_type = response.headers.get('Content-Type', '')
            if 'video' not in content_type:
                return jsonify({'error': 'Unsupported video format from URL'}), 400
            video_path = f"/tmp/{url.split('/')[-1]}"
            with open(video_path, 'wb') as f:
                f.write(response.content)
        except requests.RequestException as e:
            logging.error(f"Error fetching video from URL: {e}")
            return jsonify({'error': 'Error fetching video from URL'}), 500
    else:
        return jsonify({'error': 'No file or URL provided'}), 400

    try:
        transcript = get_transcript_from_video(video_path)
        return jsonify({'transcript': transcript})
    except Exception as e:
        logging.error(f"An error occurred during transcript extraction: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True)
